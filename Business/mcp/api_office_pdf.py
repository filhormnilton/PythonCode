"""
MCP: api_office_pdf
Tools for creating, reading, updating, and deleting Word (.docx), PDF, and TXT documents.
"""
import os
from pathlib import Path
from typing import Optional

from langchain_core.tools import tool

from Business.config import CONFIG


def _ensure_dir(directory: Path) -> None:
    directory.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Word (.docx) tools
# ---------------------------------------------------------------------------

@tool
def create_word_document(filename: str, content: str) -> str:
    """Create a new Word (.docx) document with the given content.

    Args:
        filename: Name of the file (without extension).
        content: Text content to write into the document.

    Returns:
        Absolute path of the created document.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return "ERROR: python-docx is not installed. Run: pip install python-docx"

    _ensure_dir(CONFIG.output.docs_dir)
    path = CONFIG.output.docs_dir / f"{filename}.docx"
    doc = Document()
    for paragraph in content.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return str(path)


@tool
def read_word_document(filepath: str) -> str:
    """Read and return the full text of a Word (.docx) document.

    Args:
        filepath: Absolute or relative path to the .docx file.

    Returns:
        Full text content of the document.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return "ERROR: python-docx is not installed. Run: pip install python-docx"

    doc = Document(filepath)
    return "\n".join(p.text for p in doc.paragraphs)


@tool
def update_word_document(filepath: str, new_content: str) -> str:
    """Overwrite an existing Word (.docx) document with new content.

    Args:
        filepath: Absolute or relative path to the .docx file.
        new_content: Replacement text content.

    Returns:
        Confirmation message with the file path.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return "ERROR: python-docx is not installed. Run: pip install python-docx"

    doc = Document()
    for paragraph in new_content.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(filepath)
    return f"Document updated: {filepath}"


@tool
def delete_document(filepath: str) -> str:
    """Delete a document file (Word, PDF, or TXT).

    Args:
        filepath: Absolute or relative path to the file.

    Returns:
        Confirmation message.
    """
    path = Path(filepath)
    if not path.exists():
        return f"ERROR: File not found: {filepath}"
    path.unlink()
    return f"Deleted: {filepath}"


# ---------------------------------------------------------------------------
# PDF tools
# ---------------------------------------------------------------------------

@tool
def create_pdf_document(filename: str, content: str) -> str:
    """Create a new PDF document with the given text content.

    Args:
        filename: Name of the file (without extension).
        content: Text content to render into the PDF.

    Returns:
        Absolute path of the created PDF.
    """
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.platypus import Paragraph, SimpleDocTemplate  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
    except ImportError:
        return "ERROR: reportlab is not installed. Run: pip install reportlab"

    _ensure_dir(CONFIG.output.docs_dir)
    path = CONFIG.output.docs_dir / f"{filename}.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(line.replace("\n", "<br/>"), styles["Normal"]) for line in content.split("\n\n")]
    doc.build(story)
    return str(path)


@tool
def read_pdf_document(filepath: str) -> str:
    """Extract and return text from a PDF document.

    Args:
        filepath: Absolute or relative path to the .pdf file.

    Returns:
        Extracted text content.
    """
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return "ERROR: pdfplumber is not installed. Run: pip install pdfplumber"

    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


# ---------------------------------------------------------------------------
# TXT tools
# ---------------------------------------------------------------------------

@tool
def create_txt_document(filename: str, content: str) -> str:
    """Create a plain-text (.txt) file with the given content.

    Args:
        filename: Name of the file (without extension).
        content: Text content.

    Returns:
        Absolute path of the created file.
    """
    _ensure_dir(CONFIG.output.docs_dir)
    path = CONFIG.output.docs_dir / f"{filename}.txt"
    path.write_text(content, encoding="utf-8")
    return str(path)


@tool
def read_txt_document(filepath: str) -> str:
    """Read a plain-text (.txt) file.

    Args:
        filepath: Absolute or relative path to the .txt file.

    Returns:
        File content as a string.
    """
    return Path(filepath).read_text(encoding="utf-8")


@tool
def update_txt_document(filepath: str, new_content: str) -> str:
    """Overwrite a plain-text file with new content.

    Args:
        filepath: Absolute or relative path to the .txt file.
        new_content: Replacement content.

    Returns:
        Confirmation message.
    """
    Path(filepath).write_text(new_content, encoding="utf-8")
    return f"TXT updated: {filepath}"


# ---------------------------------------------------------------------------
# Exported tool list
# ---------------------------------------------------------------------------

OFFICE_PDF_TOOLS = [
    create_word_document,
    read_word_document,
    update_word_document,
    delete_document,
    create_pdf_document,
    read_pdf_document,
    create_txt_document,
    read_txt_document,
    update_txt_document,
]
